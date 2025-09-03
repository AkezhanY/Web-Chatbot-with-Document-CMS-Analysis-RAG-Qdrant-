# backend/document_parser.py
from __future__ import annotations
import os
from typing import List, Tuple

import pandas as pd

# --- PDF: pdfminer.six ---
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None

# --- DOCX: python-docx ---
try:
    from docx import Document
except Exception:
    Document = None


def _clean(s: str) -> str:
    return " ".join((s or "").split())


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
    text = _clean(text)
    if not text:
        return []
    out, i = [], 0
    step = max(1, chunk_size - overlap)
    while i < len(text):
        out.append(text[i:i + chunk_size])
        i += step
    return out


# ---------- Readers ----------
def read_txt(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "utf-16", "cp1251", "latin-1"):
        try:
            with open(path, "r", encoding=enc, errors="ignore") as f:
                return _clean(f.read())
        except Exception:
            continue
    return ""


def read_csv(path: str) -> str:
    for sep in (",", ";", "\t", "|"):
        try:
            df = pd.read_csv(path, sep=sep)
            return df.to_csv(index=False)
        except Exception:
            continue
    try:
        df = pd.read_csv(path)
        return df.to_csv(index=False)
    except Exception:
        return ""


def read_pdf(path: str) -> str:
    if pdf_extract_text is None:
        return ""
    try:
        txt = pdf_extract_text(path) or ""
        return _clean(txt)
    except Exception:
        return ""


def read_docx(path: str) -> str:
    if Document is None:
        return ""
    try:
        d = Document(path)
        parts: List[str] = []
        parts.extend(p.text for p in d.paragraphs)
        for tbl in d.tables:
            for row in tbl.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return _clean("\n".join(parts))
    except Exception:
        return ""


# ---------- RAW fallback  ----------
def _raw_text(path: str) -> str:
    """
    Считывает байты и пытается выжать текст даже при «кривой» кодировке.
    Это спасает .txt/.md/.json и многие .csv.
    """
    try:
        with open(path, "rb") as f:
            b = f.read()
        
        return _clean(b.decode("utf-8", "ignore"))
    except Exception:
        return ""


def read_file(path: str) -> tuple[str, str]:
    ext = os.path.splitext(path)[1].lower()

    # основная попытка по типу
    if ext in (".txt", ".md", ".json"):
        text, ftype = read_txt(path), "txt"
    elif ext == ".csv":
        text, ftype = read_csv(path), "csv"
    elif ext == ".pdf":
        text, ftype = read_pdf(path), "pdf"
    elif ext == ".docx":
        text, ftype = read_docx(path), "docx"
    else:
        text, ftype = read_txt(path), "txt"

    
    if not (text or "").strip():
        
        if ext in (".txt", ".md", ".json", ".csv"):
            text = _raw_text(path)

    return text, ftype


def parse_to_chunks(path: str) -> Tuple[List[str], str, str]:
    """
    Возвращает (chunks, file_type, full_text).
    """
    text, ftype = read_file(path)
    chunks = chunk_text(text) if text else []
    return chunks, ftype, text or ""

